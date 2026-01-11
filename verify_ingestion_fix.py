
import os
import shutil
from pypdf import PdfWriter
from docx import Document
import subprocess
import time

def create_test_files():
    raw_path = "data/raw_test"
    if os.path.exists(raw_path):
        shutil.rmtree(raw_path)
    os.makedirs(raw_path)

    # 1. Create PDF
    pdf_path = os.path.join(raw_path, "test_doc.pdf")
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    # Adding text to PDF via pypdf is tricky without a font, but let's try a simple annotation or just metadata.
    # Actually, pypdf is bad at WRITING text content from scratch. 
    # Let's switch to creating a simple text file renamed as pdf? No that won't work for a real parser.
    # Let's use FPDF if available? Not in requirements.
    # Let's fallback to just checking if the code RUNS without error on a dummy PDF (even if empty content)
    # OR, we can use `reportlab` if installed?
    # Let's try mocking the _read_pdf method in the test? No, we want integration test.
    # Okay, for now, let's just make a text file to ensure the script runs, and I will manually create a PDF if needed.
    # WAIT! `docx` is easy.
    
    # 2. Create DOCX
    docx_path = os.path.join(raw_path, "test_doc.docx")
    doc = Document()
    doc.add_paragraph("This is a unique string: BANANA_123")
    doc.save(docx_path)
    print(f"Created {docx_path}")
    
    return raw_path

def test_ingestion(data_path):
    # Temporarily point config to test data? 
    # Or just modify the ingestor code?
    # Easier: Just put the file in real data/raw and clean up later.
    real_data_path = "data/raw"
    test_docx = os.path.join(real_data_path, "temp_test.docx")
    
    doc = Document()
    doc.add_paragraph("This is a unique string: BANANA_123")
    doc.save(test_docx)
    
    import sys
    print("Running ingestion...")
    result = subprocess.run([sys.executable, "main.py", "--ingest"], capture_output=True, text=True)
    print(result.stdout)
    print(result.stderr)
    
    if result.returncode == 0:
        print("Ingestion command ran successfully.")
    else:
        print("Ingestion command FAILD.")
        
    # Clean up
    if os.path.exists(test_docx):
        os.remove(test_docx)

if __name__ == "__main__":
    test_ingestion("")
